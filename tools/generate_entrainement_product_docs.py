from __future__ import annotations

from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "generated"


ENTRAINEMENT_DOC = OUTPUT_DIR / "Sport_Insight_Entrainement_Sprint2.docx"
ENTRAINEMENT_SEQUENCE = OUTPUT_DIR / "entrainement_sequence.png"
ENTRAINEMENT_LOGICAL = OUTPUT_DIR / "entrainement_logical_architecture.png"
ENTRAINEMENT_PHYSICAL = OUTPUT_DIR / "entrainement_physical_architecture.png"

PRODUIT_DOC = OUTPUT_DIR / "Sport_Insight_Produit_Sprint2.docx"
PRODUIT_SEQUENCE = OUTPUT_DIR / "produit_sequence.png"
PRODUIT_LOGICAL = OUTPUT_DIR / "produit_logical_architecture.png"
PRODUIT_PHYSICAL = OUTPUT_DIR / "produit_physical_architecture.png"


ENTRAINEMENT_BACKLOG = [
    {
        "id": "US-EN-01",
        "story": "En tant que coach, je souhaite creer et modifier une session d'entrainement avec date, horaire, type, objectif et lieu.",
        "priority": "HAUTE",
        "points": "5",
        "tasks": "Formulaire JavaFX, validation date/horaires, CRUD via EntrainementService, table admin et cartes utilisateur",
        "owner": "Tesnim",
        "status": "Termine",
    },
    {
        "id": "US-EN-02",
        "story": "En tant qu'utilisateur, je souhaite rechercher et trier les entrainements pour trouver rapidement une session.",
        "priority": "HAUTE",
        "points": "3",
        "tasks": "SearchField, sortChoiceBox, filtres sur type/lieu/objectif, mise a jour dynamique des cartes",
        "owner": "Tesnim",
        "status": "Termine",
    },
    {
        "id": "US-EN-03",
        "story": "En tant que coach, je souhaite enregistrer la presence ou l'absence pour une session d'entrainement.",
        "priority": "HAUTE",
        "points": "5",
        "tasks": "ParticipationService, save/update participation, justification d'absence, contexte de session selectionnee",
        "owner": "Tesnim",
        "status": "Termine",
    },
    {
        "id": "US-EN-04",
        "story": "En tant que coach, je souhaite gerer les evaluations physique, technique et tactique des joueurs.",
        "priority": "HAUTE",
        "points": "5",
        "tasks": "EvaluationService, formulaire d'evaluation, moyenne, tri par notes, table admin des evaluations",
        "owner": "Tesnim",
        "status": "Termine",
    },
    {
        "id": "US-EN-05",
        "story": "En tant qu'utilisateur, je souhaite visualiser les evaluations de la session selectionnee dans l'interface entrainement.",
        "priority": "MOYENNE",
        "points": "3",
        "tasks": "refreshParticipationAndEvaluations(), cartes d'evaluation, selection d'une session depuis les cards",
        "owner": "Tesnim",
        "status": "Termine",
    },
    {
        "id": "US-EN-06",
        "story": "En tant qu'administrateur, je souhaite disposer d'un dashboard avec les onglets Entrainements, Evaluations et Participations.",
        "priority": "MOYENNE",
        "points": "3",
        "tasks": "TabPane admin, TableView par sous-module, actions d'actualisation, suppression et reinitialisation",
        "owner": "Tesnim",
        "status": "Termine",
    },
]

ENTRAINEMENT_STORY = (
    "En tant que coach, je souhaite enregistrer ma participation a une session d'entrainement "
    "et consulter les evaluations liees a cette session."
)

ENTRAINEMENT_ACCEPTANCE = [
    "Etant donne un coach connecte, une session d'entrainement selectionnee et les services ParticipationService et EvaluationService disponibles",
    "Quand le coach choisit Present ou Absent puis enregistre sa participation",
    "Alors la participation est ajoutee ou mise a jour en base et les evaluations associees a la session sont rechargees dans l'interface.",
]

ENTRAINEMENT_REJECTION = [
    "Etant donne aucune session selectionnee, aucun coach connecte ou un champ presence non renseigne",
    "Quand l'utilisateur tente d'enregistrer la participation",
    "Alors un message de validation s'affiche et aucune participation n'est enregistree.",
]

ENTRAINEMENT_LOGICAL_BULLETS = [
    "Couche presentation : EntrainementAdminController et EntrainementUserController pilotent les vues admin et utilisateur.",
    "Couche metier : EntrainementService, EvaluationService et ParticipationService encapsulent le CRUD sur les trois sous-domaines.",
    "Couche securite : AuthSession et UserRoles limitent les actions de gestion aux comptes coach.",
    "Couche donnees : entites Entrainement, Evaluation et Participation echanges avec MySQL via MyConnection.",
]

ENTRAINEMENT_PHYSICAL_BULLETS = [
    "Poste client : application desktop JavaFX executee localement.",
    "Module entrainement : ecrans utilisateur et admin dans le meme client.",
    "Base locale : MySQL sport_insight sur 127.0.0.1:3306.",
    "Protocoles : interaction UI locale et JDBC/SQL pour la persistance.",
]


PRODUIT_BACKLOG = [
    {
        "id": "US-PR-01",
        "story": "En tant qu'administrateur, je souhaite ajouter, modifier et supprimer un produit dans le catalogue.",
        "priority": "HAUTE",
        "points": "5",
        "tasks": "Formulaire produit, ProductService, validations nom/categorie/prix/stock, actions table et detail panel",
        "owner": "Amine",
        "status": "Termine",
    },
    {
        "id": "US-PR-02",
        "story": "En tant qu'administrateur, je souhaite rechercher et trier les produits pour analyser rapidement le catalogue.",
        "priority": "HAUTE",
        "points": "3",
        "tasks": "searchField, ProductRepository.ProductSortField, ordre ASC/DESC, refreshProducts()",
        "owner": "Amine",
        "status": "Termine",
    },
    {
        "id": "US-PR-03",
        "story": "En tant qu'administrateur, je souhaite visualiser le stock et la repartition des categories.",
        "priority": "MOYENNE",
        "points": "3",
        "tasks": "BarChart stockStatusChart, PieChart categoryDistributionChart, metriques low stock et rupture",
        "owner": "Amine",
        "status": "Termine",
    },
    {
        "id": "US-PR-04",
        "story": "En tant qu'administrateur, je souhaite exporter la liste des produits en PDF.",
        "priority": "MOYENNE",
        "points": "2",
        "tasks": "ProductPdfExportService, choix du fichier cible, export PDF avec resume du stock",
        "owner": "Amine",
        "status": "Termine",
    },
    {
        "id": "US-PR-05",
        "story": "En tant qu'utilisateur, je souhaite ajouter des produits au panier et modifier les quantites.",
        "priority": "HAUTE",
        "points": "5",
        "tasks": "StoreController, CartLine, catalogue, panier, total dynamique, synchronisation avec le stock",
        "owner": "Amine",
        "status": "Termine",
    },
    {
        "id": "US-PR-06",
        "story": "En tant qu'utilisateur, je souhaite valider le paiement pour enregistrer la commande et recevoir une facture PDF.",
        "priority": "HAUTE",
        "points": "5",
        "tasks": "OrderService, reduction de stock, validation checkout, OrderPdfExportService, facture dans Downloads",
        "owner": "Amine",
        "status": "Termine",
    },
]

PRODUIT_STORY = (
    "En tant qu'utilisateur, je souhaite payer mon panier depuis le store afin d'enregistrer la commande, "
    "mettre a jour le stock et generer automatiquement une facture PDF."
)

PRODUIT_ACCEPTANCE = [
    "Etant donne un utilisateur connecte, un panier non vide, un produit en stock et des informations de paiement et de livraison valides",
    "Quand l'utilisateur clique sur Pay now",
    "Alors la commande est enregistree, le stock des produits est mis a jour, le panier est vide et une facture PDF est generee automatiquement.",
]

PRODUIT_REJECTION = [
    "Etant donne un panier vide, un formulaire de paiement invalide ou un stock insuffisant",
    "Quand l'utilisateur tente de finaliser la commande",
    "Alors un message de validation ou d'erreur s'affiche et aucune commande n'est enregistree.",
]

PRODUIT_LOGICAL_BULLETS = [
    "Couche presentation : ProductController pour l'admin et StoreController pour l'experience utilisateur.",
    "Couche metier : ProductService gere le catalogue ; OrderService gere les commandes et le decrement du stock.",
    "Couche export : ProductPdfExportService et OrderPdfExportService generent les documents PDF.",
    "Couche donnees : Product, Order et MyConnection s'appuient sur MySQL et le depot ProductRepository pour le tri.",
]

PRODUIT_PHYSICAL_BULLETS = [
    "Poste client : application JavaFX locale avec deux interfaces, admin produit et store utilisateur.",
    "Base locale : MySQL sport_insight sur 127.0.0.1:3306 pour les tables product et order.",
    "Stockage local : generation des PDF produits et factures sur le poste de l'utilisateur, notamment dans Downloads pour les factures.",
    "Protocoles : JDBC/SQL vers MySQL et ecriture de fichiers PDF locale, sans passerelle de paiement externe.",
]


def ensure_output_dir() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = []
    if bold:
        candidates += [
            r"C:\Windows\Fonts\arialbd.ttf",
            r"C:\Windows\Fonts\calibrib.ttf",
            r"C:\Windows\Fonts\segoeuib.ttf",
        ]
    else:
        candidates += [
            r"C:\Windows\Fonts\arial.ttf",
            r"C:\Windows\Fonts\calibri.ttf",
            r"C:\Windows\Fonts\segoeui.ttf",
        ]

    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    words = text.split()
    if not words:
        return [""]

    lines: list[str] = []
    current = words[0]
    for word in words[1:]:
        trial = f"{current} {word}"
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
            current = trial
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def draw_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    fill: str,
    outline: str,
    font: ImageFont.ImageFont,
    text_fill: str = "#202124",
    radius: int = 16,
) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, (x2 - x1) - 24)
    total_height = len(lines) * 28
    y = y1 + ((y2 - y1) - total_height) // 2 - 3
    for line in lines:
        width = draw.textbbox((0, 0), line, font=font)[2]
        x = x1 + ((x2 - x1) - width) // 2
        draw.text((x, y), line, font=font, fill=text_fill)
        y += 28


def draw_lifeline(draw: ImageDraw.ImageDraw, x: int, top: int, bottom: int, color: str = "#2f2e4a") -> None:
    draw.line((x, top, x, bottom), fill=color, width=2)


def draw_arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    text: str,
    font: ImageFont.ImageFont,
    number: int | None = None,
    dashed: bool = False,
    text_offset: int = -24,
    color: str = "#23233f",
) -> None:
    x1, y1 = start
    x2, y2 = end

    if dashed:
        dash = 10
        gap = 6
        total = abs(x2 - x1)
        direction = 1 if x2 >= x1 else -1
        current = x1
        while abs(current - x1) < total:
            next_x = current + direction * min(dash, total - abs(current - x1))
            draw.line((current, y1, next_x, y2), fill=color, width=3)
            current = next_x + direction * gap
    else:
        draw.line((x1, y1, x2, y2), fill=color, width=3)

    arrow_size = 12
    if x2 >= x1:
        arrow = [(x2, y2), (x2 - arrow_size, y2 - 7), (x2 - arrow_size, y2 + 7)]
    else:
        arrow = [(x2, y2), (x2 + arrow_size, y2 - 7), (x2 + arrow_size, y2 + 7)]
    draw.polygon(arrow, fill=color)

    mid_x = (x1 + x2) // 2
    lines = wrap_text(draw, text, font, max(150, abs(x2 - x1) - 30))
    base_y = y1 + text_offset
    for idx, line in enumerate(lines):
        width = draw.textbbox((0, 0), line, font=font)[2]
        draw.text((mid_x - width // 2, base_y + idx * 22), line, font=font, fill="#111111")

    if number is not None:
        circle_x = x1 + 16 if x2 >= x1 else x1 - 16
        circle_y = y1
        draw.ellipse((circle_x - 10, circle_y - 10, circle_x + 10, circle_y + 10), fill="#2f2e4a")
        num = str(number)
        width = draw.textbbox((0, 0), num, font=font)[2]
        draw.text((circle_x - width // 2, circle_y - 9), num, font=font, fill="white")


def draw_sequence_diagram(
    path: Path,
    title: str,
    participants: list[tuple[str, str, str]],
    x_positions: list[int],
    events: list[tuple[int, int, str, bool]],
) -> None:
    width = 1900
    height = max(1120, 360 + len(events) * 64 + 180)
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    title_font = get_font(28, bold=True)
    box_font = get_font(20, bold=False)
    arrow_font = get_font(18, bold=False)

    draw.text((width // 2 - draw.textbbox((0, 0), title, font=title_font)[2] // 2, 24), title, font=title_font, fill="#1f1f1f")

    top_box_y = 88
    bottom_y = height - 100
    for (label, fill, outline), x in zip(participants, x_positions):
        draw_box(draw, (x - 95, top_box_y, x + 95, top_box_y + 74), label, fill, outline, box_font)
        draw_lifeline(draw, x, top_box_y + 86, bottom_y)
        draw_box(draw, (x - 95, bottom_y, x + 95, bottom_y + 74), label, fill, outline, box_font)

    y = 210
    for idx, (source, target, label, dashed) in enumerate(events, start=1):
        draw_arrow(
            draw,
            (x_positions[source], y),
            (x_positions[target], y),
            label,
            arrow_font,
            number=idx,
            dashed=dashed,
            text_offset=-24 if not dashed else -18,
        )
        y += 62

    image.save(path)


def draw_layered_architecture(
    path: Path,
    title: str,
    layers: list[tuple[str, str, str, list[str]]],
    footer_bullets: list[str],
) -> None:
    width, height = 1700, 980
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = get_font(28, bold=True)
    box_font = get_font(22, bold=False)
    small_font = get_font(18, bold=False)

    draw.text((width // 2 - draw.textbbox((0, 0), title, font=title_font)[2] // 2, 26), title, font=title_font, fill="#222222")

    x1, x2 = 90, width - 90
    y = 110
    layer_height = 170
    gap = 24

    for index, (layer_name, fill, outline, items) in enumerate(layers):
        top = y + index * (layer_height + gap)
        bottom = top + layer_height
        draw.rounded_rectangle((x1, top, x2, bottom), radius=22, fill=fill, outline=outline, width=4)
        draw.text((x1 + 28, top + 18), layer_name, font=box_font, fill="#1e1e1e")

        item_y = top + 64
        item_gap = (x2 - x1 - 80) // len(items)
        for item_index, item in enumerate(items):
            left = x1 + 28 + item_index * item_gap
            right = left + item_gap - 18
            draw.rounded_rectangle((left, item_y, right, item_y + 70), radius=16, fill="white", outline=outline, width=3)
            lines = wrap_text(draw, item, small_font, right - left - 18)
            text_y = item_y + 16
            for line in lines:
                line_width = draw.textbbox((0, 0), line, font=small_font)[2]
                draw.text((left + (right - left - line_width) // 2, text_y), line, font=small_font, fill="#222222")
                text_y += 22

        if index < len(layers) - 1:
            mid_x = width // 2
            draw.line((mid_x, bottom, mid_x, bottom + gap), fill="#6b7280", width=4)
            draw.polygon([(mid_x, bottom + gap), (mid_x - 10, bottom + gap - 14), (mid_x + 10, bottom + gap - 14)], fill="#6b7280")

    footer_y = 895
    for bullet in footer_bullets:
        draw.text((90, footer_y), f"- {bullet}", font=small_font, fill="#303030")
        footer_y += 24

    image.save(path)


def draw_entrainement_physical(path: Path) -> None:
    width, height = 1700, 980
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = get_font(28, bold=True)
    box_font = get_font(22, bold=False)
    small_font = get_font(18, bold=False)

    title = "Architecture physique - Deploiement du module Entrainement"
    draw.text((width // 2 - draw.textbbox((0, 0), title, font=title_font)[2] // 2, 28), title, font=title_font, fill="#222222")

    coach_box = (110, 180, 470, 470)
    client_box = (610, 160, 1080, 520)
    db_box = (1180, 240, 1570, 560)

    draw.rounded_rectangle(coach_box, radius=24, fill="#e8f4ff", outline="#2d9cdb", width=4)
    draw.rounded_rectangle(client_box, radius=24, fill="#fff4e5", outline="#ff9800", width=4)
    draw.rounded_rectangle(db_box, radius=24, fill="#eefbe9", outline="#4caf50", width=4)

    draw.text((180, 220), "Coach / Admin", font=box_font, fill="#1f1f1f")
    for idx, line in enumerate([
        "Utilisateur connecte",
        "Selection de session",
        "Saisie presence / evaluation",
        "Interface locale JavaFX",
    ]):
        draw.text((140, 280 + idx * 32), f"- {line}", font=small_font, fill="#303030")

    draw.text((705, 210), "Client Entrainement JavaFX", font=box_font, fill="#1f1f1f")
    for idx, line in enumerate([
        "EntrainementUserController",
        "EntrainementAdminController",
        "EntrainementService",
        "EvaluationService / ParticipationService",
        "AuthSession et UserRoles",
    ]):
        draw.text((660, 270 + idx * 32), f"- {line}", font=small_font, fill="#303030")

    draw.text((1270, 285), "MySQL local", font=box_font, fill="#1f1f1f")
    for idx, line in enumerate([
        "Base sport_insight",
        "tables entrainement",
        "evaluation",
        "participation",
    ]):
        draw.text((1230, 345 + idx * 32), f"- {line}", font=small_font, fill="#303030")

    draw_arrow(draw, (470, 320), (610, 320), "Actions UI / formulaires", small_font, number=None, text_offset=-22, color="#2563eb")
    draw_arrow(draw, (1080, 360), (1180, 360), "JDBC / SQL", small_font, number=None, text_offset=-22, color="#2e7d32")

    footer_y = 690
    for bullet in ENTRAINEMENT_PHYSICAL_BULLETS:
        draw.text((120, footer_y), f"- {bullet}", font=small_font, fill="#303030")
        footer_y += 30

    image.save(path)


def draw_produit_physical(path: Path) -> None:
    width, height = 1700, 980
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = get_font(28, bold=True)
    box_font = get_font(22, bold=False)
    small_font = get_font(18, bold=False)

    title = "Architecture physique - Deploiement du module Produit / Store"
    draw.text((width // 2 - draw.textbbox((0, 0), title, font=title_font)[2] // 2, 28), title, font=title_font, fill="#222222")

    user_box = (90, 160, 430, 470)
    client_box = (520, 140, 980, 520)
    db_box = (1090, 170, 1460, 470)
    file_box = (1190, 600, 1560, 820)

    draw.rounded_rectangle(user_box, radius=24, fill="#e8f4ff", outline="#2d9cdb", width=4)
    draw.rounded_rectangle(client_box, radius=24, fill="#fff4e5", outline="#ff9800", width=4)
    draw.rounded_rectangle(db_box, radius=24, fill="#eefbe9", outline="#4caf50", width=4)
    draw.rounded_rectangle(file_box, radius=24, fill="#f3ebff", outline="#7b61ff", width=4)

    draw.text((175, 205), "Utilisateur / Admin", font=box_font, fill="#1f1f1f")
    for idx, line in enumerate([
        "Catalogue",
        "Panier",
        "Paiement",
        "CRUD admin",
    ]):
        draw.text((135, 270 + idx * 32), f"- {line}", font=small_font, fill="#303030")

    draw.text((640, 190), "Client JavaFX Produit", font=box_font, fill="#1f1f1f")
    for idx, line in enumerate([
        "ProductController",
        "StoreController",
        "ProductService",
        "OrderService",
        "PDF export services",
    ]):
        draw.text((575, 250 + idx * 32), f"- {line}", font=small_font, fill="#303030")

    draw.text((1185, 220), "MySQL local", font=box_font, fill="#1f1f1f")
    for idx, line in enumerate([
        "Base sport_insight",
        "table product",
        "table order",
    ]):
        draw.text((1150, 285 + idx * 34), f"- {line}", font=small_font, fill="#303030")

    draw.text((1260, 650), "Fichiers PDF locaux", font=box_font, fill="#1f1f1f")
    for idx, line in enumerate([
        "Export catalogue",
        "Factures client",
        "Downloads / disque local",
    ]):
        draw.text((1230, 710 + idx * 34), f"- {line}", font=small_font, fill="#303030")

    draw_arrow(draw, (430, 300), (520, 300), "Actions UI", small_font, number=None, text_offset=-22, color="#2563eb")
    draw_arrow(draw, (980, 340), (1090, 340), "JDBC / SQL", small_font, number=None, text_offset=-22, color="#2e7d32")
    draw_arrow(draw, (980, 430), (1190, 690), "Generation PDF", small_font, number=None, text_offset=-18, color="#6d28d9")

    footer_y = 860
    for bullet in PRODUIT_PHYSICAL_BULLETS:
        draw.text((95, footer_y), f"- {bullet}", font=small_font, fill="#303030")
        footer_y += 24

    image.save(path)


def set_page_margins(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.905)
    section.bottom_margin = Cm(1.905)
    section.left_margin = Cm(1.905)
    section.right_margin = Cm(1.905)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 11) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def add_highlighted_gherkin(cell, steps: Iterable[str]) -> None:
    keywords = ("Etant donne", "Quand", "Alors")
    cell.text = ""
    for index, step in enumerate(steps):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        for keyword in keywords:
            if step.startswith(keyword):
                keyword_run = paragraph.add_run(keyword)
                keyword_run.bold = True
                keyword_run.font.size = Pt(11)
                keyword_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                rest = step[len(keyword):]
                if rest:
                    rest_run = paragraph.add_run(rest)
                    rest_run.font.size = Pt(11)
                break
        else:
            run = paragraph.add_run(step)
            run.font.size = Pt(11)


def style_table_borders(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.first_child_found_in("w:tcBorders")
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for edge in ("top", "left", "bottom", "right"):
                element = borders.find(qn(f"w:{edge}"))
                if element is None:
                    element = OxmlElement(f"w:{edge}")
                    borders.append(element)
                element.set(qn("w:val"), "single")
                element.set(qn("w:sz"), "8")
                element.set(qn("w:color"), "000000")


def add_title(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(32)


def add_section_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(13)


def add_backlog_table(document: Document, backlog: list[dict[str, str]]) -> None:
    add_section_heading(document, "1. Sprint Backlog  Sprint 2")
    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["ID", "User Story", "Priorite", "Points", "Taches", "Responsable", "Statut"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "D9D9D9")

    for item in backlog:
        row = table.add_row().cells
        values = [item["id"], item["story"], item["priority"], item["points"], item["tasks"], item["owner"], item["status"]]
        for idx, value in enumerate(values):
            set_cell_text(row[idx], value)

    style_table_borders(table)
    document.add_paragraph("")


def add_story_tests(
    document: Document,
    story_label: str,
    story_under_test: str,
    acceptance_steps: list[str],
    rejection_steps: list[str],
) -> None:
    add_section_heading(document, "2. Tableau des Story Tests")
    intro = document.add_paragraph()
    run = intro.add_run("User Story testee : ")
    run.bold = True
    intro.add_run(story_label)

    table = document.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["User Story", "Story test d'acceptation", "Story Test de refus"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "D9D9D9")

    story_cell = table.rows[1].cells[0]
    story_cell.text = ""
    paragraph = story_cell.paragraphs[0]
    run = paragraph.add_run(story_under_test)
    run.bold = True
    run.italic = True
    run.font.size = Pt(11)

    add_highlighted_gherkin(table.rows[1].cells[1], acceptance_steps)
    add_highlighted_gherkin(table.rows[1].cells[2], rejection_steps)

    style_table_borders(table)
    document.add_paragraph("")


def add_image(document: Document, path: Path, width: float, caption: str | None = None) -> None:
    if caption:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(caption)
        run.bold = True
    document.add_picture(str(path), width=Inches(width))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def generate_entrainement_assets() -> None:
    participants = [
        ("Coach", "#fde7ff", "#d248f7"),
        ("EntrainementUserController", "#e7fcff", "#00b7d7"),
        ("ParticipationService", "#fff1e3", "#ff8c2a"),
        ("EvaluationService", "#e8f4ff", "#27a3ff"),
        ("MySQL", "#eafbea", "#31bf5b"),
    ]
    x_positions = [120, 480, 860, 1220, 1600]
    events = [
        (0, 1, "Selectionne une session d'entrainement", False),
        (1, 2, "findCurrentUserParticipation()", False),
        (2, 4, "SELECT participation par session/joueur", False),
        (4, 2, "Participation existante ou vide", True),
        (1, 3, "getByEntrainement(sessionId)", False),
        (3, 4, "SELECT evaluations par session", False),
        (4, 3, "Liste des evaluations", True),
        (3, 1, "Evaluations a afficher", True),
        (0, 1, "Clique sur Enregistrer participation", False),
        (1, 2, "add() ou update()", False),
        (2, 4, "INSERT / UPDATE participation", False),
        (4, 2, "Confirmation SQL", True),
        (2, 1, "Participation sauvegardee", True),
        (1, 3, "Recharger evaluations de la session", False),
        (3, 4, "SELECT evaluations actualisees", False),
        (4, 3, "Resultats", True),
        (3, 1, "Cartes d'evaluation rechargees", True),
        (1, 0, "Participation enregistree + vue mise a jour", True),
    ]
    draw_sequence_diagram(
        ENTRAINEMENT_SEQUENCE,
        "Diagramme de sequence - Participation et evaluations d'une session",
        participants,
        x_positions,
        events,
    )

    draw_layered_architecture(
        ENTRAINEMENT_LOGICAL,
        "Architecture logique - Module Entrainement",
        [
            ("Couche presentation", "#e7f4ff", "#2d9cdb", [
                "EntrainementAdminController",
                "EntrainementUserController",
                "entrainement-admin-view.fxml",
                "entrainement-user-view.fxml",
            ]),
            ("Couche metier", "#eefbe9", "#4caf50", [
                "EntrainementService",
                "EvaluationService",
                "ParticipationService",
            ]),
            ("Couche securite et support", "#fff4e5", "#ff9800", [
                "AuthSession",
                "UserRoles",
                "UserService",
            ]),
            ("Couche donnees", "#f2edff", "#7b61ff", [
                "Entrainement",
                "Evaluation",
                "Participation",
                "MyConnection / MySQL",
            ]),
        ],
        ENTRAINEMENT_LOGICAL_BULLETS,
    )

    draw_entrainement_physical(ENTRAINEMENT_PHYSICAL)


def generate_produit_assets() -> None:
    participants = [
        ("Utilisateur", "#fde7ff", "#d248f7"),
        ("StoreController", "#e7fcff", "#00b7d7"),
        ("ProductService", "#fff1e3", "#ff8c2a"),
        ("OrderService", "#e8f4ff", "#27a3ff"),
        ("MySQL", "#eafbea", "#31bf5b"),
        ("OrderPdfExportService", "#fff6df", "#e2a800"),
        ("Downloads", "#f2edff", "#7b61ff"),
    ]
    x_positions = [110, 380, 660, 940, 1220, 1500, 1760]
    events = [
        (0, 1, "Ajoute des produits au panier puis clique sur Pay now", False),
        (1, 2, "getById(productId) pour verifier le stock", False),
        (2, 4, "SELECT product FOR UPDATE", False),
        (4, 2, "Produit et stock courant", True),
        (2, 1, "Produit valide pour checkout", True),
        (1, 3, "add(order) pour chaque ligne du panier", False),
        (3, 4, "INSERT INTO order", False),
        (4, 3, "Commande inseree", True),
        (3, 4, "UPDATE product SET stock = stock - quantity", False),
        (4, 3, "Stock mis a jour", True),
        (3, 1, "Commande enregistree", True),
        (1, 5, "exportInvoiceAutomatically()", False),
        (5, 6, "Ecrit facture PDF dans Downloads", False),
        (6, 5, "PDF cree", True),
        (5, 1, "Chemin de la facture", True),
        (1, 0, "Confirmation + panier vide + facture disponible", True),
    ]
    draw_sequence_diagram(
        PRODUIT_SEQUENCE,
        "Diagramme de sequence - Paiement du panier et generation de facture",
        participants,
        x_positions,
        events,
    )

    draw_layered_architecture(
        PRODUIT_LOGICAL,
        "Architecture logique - Module Produit / Store",
        [
            ("Couche presentation", "#e7f4ff", "#2d9cdb", [
                "ProductController",
                "StoreController",
                "product-crud-view.fxml",
                "store-view.fxml",
            ]),
            ("Couche metier", "#eefbe9", "#4caf50", [
                "ProductService",
                "OrderService",
                "ProductRepository",
            ]),
            ("Couche export et support", "#fff4e5", "#ff9800", [
                "ProductPdfExportService",
                "OrderPdfExportService",
                "ProductImageResolver",
            ]),
            ("Couche donnees", "#f2edff", "#7b61ff", [
                "Product",
                "Order",
                "MyConnection / MySQL",
                "Downloads / PDFs",
            ]),
        ],
        PRODUIT_LOGICAL_BULLETS,
    )

    draw_produit_physical(PRODUIT_PHYSICAL)


def build_document(
    output_path: Path,
    title: str,
    backlog: list[dict[str, str]],
    story_label: str,
    story_text: str,
    acceptance_steps: list[str],
    rejection_steps: list[str],
    sequence_heading: str,
    sequence_image: Path,
    logical_image: Path,
    physical_image: Path,
) -> None:
    document = Document()
    set_page_margins(document)
    add_title(document, title)
    document.add_paragraph("")
    add_backlog_table(document, backlog)
    add_story_tests(document, story_label, story_text, acceptance_steps, rejection_steps)
    add_section_heading(document, sequence_heading)
    add_image(document, sequence_image, 7.0)
    document.add_paragraph("")
    add_section_heading(document, "4. Architectures logique et physique")
    add_image(document, logical_image, 6.9, "Architecture logique :")
    document.add_paragraph("")
    add_image(document, physical_image, 6.9, "Architecture physique :")
    document.save(output_path)


def main() -> None:
    ensure_output_dir()
    generate_entrainement_assets()
    generate_produit_assets()

    build_document(
        ENTRAINEMENT_DOC,
        "SPORT INSIGHT - ENTRAINEMENT",
        ENTRAINEMENT_BACKLOG,
        "US-EN-03  Enregistrer la participation et consulter les evaluations",
        ENTRAINEMENT_STORY,
        ENTRAINEMENT_ACCEPTANCE,
        ENTRAINEMENT_REJECTION,
        "3. Diagramme de Sequence  Enregistrement d'une participation et consultation des evaluations",
        ENTRAINEMENT_SEQUENCE,
        ENTRAINEMENT_LOGICAL,
        ENTRAINEMENT_PHYSICAL,
    )

    build_document(
        PRODUIT_DOC,
        "SPORT INSIGHT - PRODUIT",
        PRODUIT_BACKLOG,
        "US-PR-06  Payer le panier et generer une facture PDF",
        PRODUIT_STORY,
        PRODUIT_ACCEPTANCE,
        PRODUIT_REJECTION,
        "3. Diagramme de Sequence  Paiement du panier et generation automatique de la facture",
        PRODUIT_SEQUENCE,
        PRODUIT_LOGICAL,
        PRODUIT_PHYSICAL,
    )

    print(ENTRAINEMENT_DOC)
    print(PRODUIT_DOC)


if __name__ == "__main__":
    main()
