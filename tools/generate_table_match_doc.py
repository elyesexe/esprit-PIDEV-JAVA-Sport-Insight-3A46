from __future__ import annotations

from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "generated"
OUTPUT_DOCX = OUTPUT_DIR / "Sport_Insight_Table_Match_Sprint2.docx"
SEQUENCE_PNG = OUTPUT_DIR / "table_match_sequence.png"
LOGICAL_PNG = OUTPUT_DIR / "table_match_logical_architecture.png"
PHYSICAL_PNG = OUTPUT_DIR / "table_match_physical_architecture.png"

TITLE = "SPORT INSIGHT"
RESPONSABLE = "Elyes"


SPRINT_BACKLOG = [
    {
        "id": "US-TM-01",
        "story": "En tant qu'utilisateur, je souhaite choisir un championnat afin d'ouvrir rapidement la page du classement associe.",
        "priority": "HAUTE",
        "points": "3",
        "tasks": "Cartes competitions, logos, navigation vers league-table-view.fxml, passage du competitionCode",
        "owner": RESPONSABLE,
        "status": "Termine",
    },
    {
        "id": "US-TM-02",
        "story": "En tant qu'utilisateur, je souhaite consulter le classement officiel d'un championnat pour suivre la position des clubs.",
        "priority": "HAUTE",
        "points": "5",
        "tasks": "LeagueTableController, fetchStandings(), parsing JSON, affichage du tableau complet",
        "owner": RESPONSABLE,
        "status": "Termine",
    },
    {
        "id": "US-TM-03",
        "story": "En tant qu'utilisateur, je souhaite voir la forme recente et les meilleurs buteurs pour mieux analyser la competition.",
        "priority": "HAUTE",
        "points": "5",
        "tasks": "Form chips, bloc top scorers, service ApiFootballInsightsService, rendu des statistiques",
        "owner": RESPONSABLE,
        "status": "Termine",
    },
    {
        "id": "US-TM-04",
        "story": "En tant qu'utilisateur, je souhaite actualiser le classement et les buteurs afin d'obtenir les donnees les plus recentes.",
        "priority": "MOYENNE",
        "points": "2",
        "tasks": "Bouton Actualiser, appels asynchrones, mise a jour du statusLabel et du scorersStatusLabel",
        "owner": RESPONSABLE,
        "status": "Termine",
    },
    {
        "id": "US-TM-05",
        "story": "En tant qu'utilisateur, je souhaite consulter les meta-informations d'une saison pour connaitre le nombre de clubs, la journee et la saison.",
        "priority": "MOYENNE",
        "points": "3",
        "tasks": "competitionChipLabel, clubCountChipLabel, matchdayChipLabel, seasonChipLabel, subtitle de contexte",
        "owner": RESPONSABLE,
        "status": "Termine",
    },
    {
        "id": "US-TM-06",
        "story": "En tant qu'utilisateur, je souhaite recevoir un message clair si l'API est indisponible afin de comprendre l'echec du chargement.",
        "priority": "HAUTE",
        "points": "3",
        "tasks": "Gestion d'erreurs, Alert JavaFX, status-error, prevention de faux affichages",
        "owner": RESPONSABLE,
        "status": "Termine",
    },
]

STORY_UNDER_TEST = (
    "En tant qu'utilisateur, je souhaite consulter le classement d'un championnat "
    "afin de voir la position des clubs, leurs points et leur forme recente."
)

STORY_ACCEPTANCE = [
    "Etant donne l'utilisateur se trouve sur la page 'Leagues', que la competition 'Premier League' est disponible et que football-data.org repond correctement",
    "Quand l'utilisateur clique sur la carte 'Premier League'",
    "Alors la page 'League Table' s'ouvre et affiche le classement complet avec la position, les matchs joues, les buts, la difference de buts, les points, la forme recente ainsi que le bloc des meilleurs buteurs.",
]

STORY_REJECTION = [
    "Etant donne l'utilisateur tente d'ouvrir une competition alors que football-data.org est indisponible ou renvoie une erreur",
    "Quand le chargement du classement est lance",
    "Alors aucun faux classement n'est affiche, un message d'erreur apparait et l'utilisateur peut relancer l'actualisation.",
]

LOGICAL_BULLETS = [
    "Couche presentation : LeagueCompetitionController, LeagueTableController et les vues FXML JavaFX.",
    "Couche metier : FootballDataStandingsService orchestre le chargement et la transformation du classement.",
    "Couche analyse : ApiFootballInsightsService charge les meilleurs buteurs et complete l'experience.",
    "Couche integration : FootballDataApiClient interroge football-data.org ; les DTO LeagueStandingsSnapshot et LeagueStandingEntry transportent les donnees.",
    "Couche persistance transverse : MyConnection et MySQL sont disponibles pour le module Match global de Sport Insight.",
]

PHYSICAL_BULLETS = [
    "Poste client : application desktop JavaFX executee localement sur le PC de l'utilisateur.",
    "Serveur de donnees local : MySQL 'sport_insight' sur 127.0.0.1:3306.",
    "Services externes : football-data.org pour les standings et API-Football comme source additionnelle pour les meilleurs buteurs.",
    "Protocoles : HTTP/JSON pour les APIs et JDBC pour la base MySQL.",
]


def ensure_output_dir() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = []
    if bold:
        candidates += [
            r"C:\Windows\Fonts\arialbd.ttf",
            r"C:\Windows\Fonts\calibrib.ttf",
            r"C:\Windows\Fonts\segoeuib.ttf",
        ]
    else:
        candidates += [
            r"C:\Windows\Fonts\arial.ttf",
            r"C:\Windows\Fonts\calibri.ttf",
            r"C:\Windows\Fonts\segoeui.ttf",
        ]

    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    words = text.split()
    if not words:
        return [""]

    lines: list[str] = []
    current = words[0]
    for word in words[1:]:
        trial = f"{current} {word}"
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
            current = trial
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def draw_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    fill: str,
    outline: str,
    font: ImageFont.ImageFont,
    text_fill: str = "#202124",
    radius: int = 16,
) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, (x2 - x1) - 28)
    total_height = len(lines) * 28
    y = y1 + ((y2 - y1) - total_height) // 2 - 2
    for line in lines:
        width = draw.textbbox((0, 0), line, font=font)[2]
        x = x1 + ((x2 - x1) - width) // 2
        draw.text((x, y), line, font=font, fill=text_fill)
        y += 28


def draw_arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    text: str,
    font: ImageFont.ImageFont,
    number: int | None = None,
    dashed: bool = False,
    text_offset: int = -28,
    color: str = "#23233f",
) -> None:
    x1, y1 = start
    x2, y2 = end
    if dashed:
        dash = 10
        gap = 6
        total = abs(x2 - x1)
        direction = 1 if x2 >= x1 else -1
        current = x1
        while abs(current - x1) < total:
            next_x = current + direction * min(dash, total - abs(current - x1))
            draw.line((current, y1, next_x, y2), fill=color, width=3)
            current = next_x + direction * gap
    else:
        draw.line((x1, y1, x2, y2), fill=color, width=3)

    arrow_size = 12
    if x2 >= x1:
        arrow = [(x2, y2), (x2 - arrow_size, y2 - 7), (x2 - arrow_size, y2 + 7)]
    else:
        arrow = [(x2, y2), (x2 + arrow_size, y2 - 7), (x2 + arrow_size, y2 + 7)]
    draw.polygon(arrow, fill=color)

    label = text
    mid_x = (x1 + x2) // 2
    label_lines = wrap_text(draw, label, font, max(180, abs(x2 - x1) - 36))
    label_y = y1 + text_offset
    for idx, line in enumerate(label_lines):
        width = draw.textbbox((0, 0), line, font=font)[2]
        draw.text((mid_x - width // 2, label_y + idx * 22), line, font=font, fill="#111111")

    if number is not None:
        circle_x = x1 + 16 if x2 >= x1 else x1 - 16
        circle_y = y1
        draw.ellipse((circle_x - 10, circle_y - 10, circle_x + 10, circle_y + 10), fill="#2f2e4a")
        num = str(number)
        width = draw.textbbox((0, 0), num, font=font)[2]
        draw.text((circle_x - width // 2, circle_y - 9), num, font=font, fill="white")


def draw_lifeline(draw: ImageDraw.ImageDraw, x: int, top: int, bottom: int, color: str = "#2f2e4a") -> None:
    draw.line((x, top, x, bottom), fill=color, width=2)


def create_sequence_diagram(path: Path) -> None:
    width, height = 1900, 1080
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    title_font = get_font(28, bold=True)
    label_font = get_font(22, bold=False)
    box_font = get_font(20, bold=False)
    arrow_font = get_font(18, bold=False)

    title = "Diagramme de sequence - Consultation du classement d'une competition"
    draw.text((width // 2 - draw.textbbox((0, 0), title, font=title_font)[2] // 2, 24), title, font=title_font, fill="#1f1f1f")

    participants = [
        ("Utilisateur", "#fde7ff", "#d248f7"),
        ("Page Leagues", "#e7fcff", "#00b7d7"),
        ("LeagueCompetitionController", "#fff1e3", "#ff8c2a"),
        ("LeagueTableController", "#e8f4ff", "#27a3ff"),
        ("FootballDataStandingsService", "#eafbea", "#31bf5b"),
        ("football-data.org", "#eef3ff", "#7866ff"),
        ("ApiFootballInsightsService", "#fff6df", "#e2a800"),
    ]
    x_positions = [100, 350, 610, 900, 1190, 1480, 1730]
    top_box_y = 90
    bottom_y = 990

    for (label, fill, outline), x in zip(participants, x_positions):
        draw_box(draw, (x - 90, top_box_y, x + 90, top_box_y + 74), label, fill, outline, box_font)
        draw_lifeline(draw, x, top_box_y + 86, bottom_y)
        draw_box(draw, (x - 90, bottom_y, x + 90, bottom_y + 74), label, fill, outline, box_font)

    events = [
        (0, 1, "Clique sur une carte de championnat", False),
        (1, 2, "Selection de la competition", False),
        (2, 3, "switchScene(league-table-view.fxml)", False),
        (3, 4, "loadStandingsAsync()", False),
        (4, 5, "GET /competitions/{code}/standings", False),
        (5, 4, "JSON du classement", True),
        (4, 3, "LeagueStandingsSnapshot", True),
        (3, 6, "loadCompetitionTopScorers(code)", False),
        (6, 5, "Requete scorers / fallback data", False),
        (5, 6, "Liste des meilleurs buteurs", True),
        (6, 3, "Scorers enrichis", True),
        (3, 1, "Mise a jour de l'ecran: tableau + buteurs", False),
        (1, 0, "Affiche le classement actualise", True),
    ]

    y = 210
    for idx, (source_idx, target_idx, text, dashed) in enumerate(events, start=1):
        draw_arrow(
            draw,
            (x_positions[source_idx], y),
            (x_positions[target_idx], y),
            text,
            arrow_font,
            number=idx,
            dashed=dashed,
            text_offset=-24 if not dashed else -18,
        )
        y += 65

    image.save(path)


def draw_layered_architecture(path: Path) -> None:
    width, height = 1700, 980
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = get_font(28, bold=True)
    box_font = get_font(22, bold=False)
    small_font = get_font(18, bold=False)

    title = "Architecture logique - Module Table Match / League Table"
    draw.text((width // 2 - draw.textbbox((0, 0), title, font=title_font)[2] // 2, 26), title, font=title_font, fill="#222222")

    layers = [
        ("Couche presentation", "#e7f4ff", "#2d9cdb", [
            "LeagueCompetitionController",
            "LeagueTableController",
            "league-competitions-view.fxml",
            "league-table-view.fxml",
        ]),
        ("Couche metier", "#eefbe9", "#4caf50", [
            "FootballDataStandingsService",
            "ApiFootballInsightsService",
            "FootballDataCompetitions",
        ]),
        ("Couche integration", "#fff4e5", "#ff9800", [
            "FootballDataApiClient",
            "LeagueStandingsSnapshot",
            "LeagueStandingEntry",
        ]),
        ("Sources externes et support", "#f2edff", "#7b61ff", [
            "football-data.org",
            "API-Football",
            "MyConnection / MySQL",
        ]),
    ]

    x1, x2 = 120, width - 120
    y = 110
    layer_height = 170
    gap = 26
    for index, (layer_name, fill, outline, items) in enumerate(layers):
        top = y + index * (layer_height + gap)
        bottom = top + layer_height
        draw.rounded_rectangle((x1, top, x2, bottom), radius=22, fill=fill, outline=outline, width=4)
        draw.text((x1 + 28, top + 18), layer_name, font=box_font, fill="#1e1e1e")

        item_y = top + 64
        item_gap = (x2 - x1 - 80) // len(items)
        for item_index, item in enumerate(items):
            left = x1 + 28 + item_index * item_gap
            right = left + item_gap - 20
            draw.rounded_rectangle((left, item_y, right, item_y + 70), radius=16, fill="white", outline=outline, width=3)
            lines = wrap_text(draw, item, small_font, right - left - 20)
            text_y = item_y + 16
            for line in lines:
                line_width = draw.textbbox((0, 0), line, font=small_font)[2]
                draw.text((left + (right - left - line_width) // 2, text_y), line, font=small_font, fill="#222222")
                text_y += 22

        if index < len(layers) - 1:
            mid_x = width // 2
            draw.line((mid_x, bottom, mid_x, bottom + gap), fill="#6b7280", width=4)
            draw.polygon([(mid_x, bottom + gap), (mid_x - 10, bottom + gap - 14), (mid_x + 10, bottom + gap - 14)], fill="#6b7280")

    footer_y = 900
    for bullet in LOGICAL_BULLETS:
        draw.text((120, footer_y), f"- {bullet}", font=small_font, fill="#303030")
        footer_y += 24

    image.save(path)


def draw_physical_architecture(path: Path) -> None:
    width, height = 1700, 980
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = get_font(28, bold=True)
    box_font = get_font(22, bold=False)
    small_font = get_font(18, bold=False)

    title = "Architecture physique - Deploiement du module Table Match"
    draw.text((width // 2 - draw.textbbox((0, 0), title, font=title_font)[2] // 2, 26), title, font=title_font, fill="#222222")

    client_box = (90, 160, 470, 500)
    cloud_box = (980, 110, 1570, 540)
    api1_box = (1070, 200, 1510, 320)
    api2_box = (1070, 370, 1510, 490)
    db_box = (540, 610, 980, 850)

    draw.rounded_rectangle(client_box, radius=24, fill="#e8f4ff", outline="#2d9cdb", width=4)
    draw.rounded_rectangle(cloud_box, radius=44, fill="#f5f0ff", outline="#7b61ff", width=4)
    draw.rounded_rectangle(api1_box, radius=18, fill="white", outline="#7b61ff", width=3)
    draw.rounded_rectangle(api2_box, radius=18, fill="white", outline="#7b61ff", width=3)
    draw.rounded_rectangle(db_box, radius=24, fill="#eefbe9", outline="#4caf50", width=4)

    draw.text((160, 200), "Poste utilisateur", font=box_font, fill="#1f1f1f")
    client_lines = [
        "Application desktop JavaFX",
        "LeagueCompetitionController",
        "LeagueTableController",
        "Vues FXML et composants UI",
    ]
    y = 260
    for line in client_lines:
        draw.text((120, y), f"- {line}", font=small_font, fill="#303030")
        y += 30

    draw.text((665, 650), "Serveur local MySQL", font=box_font, fill="#1f1f1f")
    db_lines = [
        "127.0.0.1:3306",
        "Base: sport_insight",
        "Acces JDBC via MyConnection",
    ]
    y = 715
    for line in db_lines:
        draw.text((590, y), f"- {line}", font=small_font, fill="#303030")
        y += 30

    draw.text((1225, 150), "Services externes", font=box_font, fill="#1f1f1f")
    draw.text((1235, 250), "football-data.org", font=box_font, fill="#1f1f1f")
    draw.text((1225, 286), "Standings officiels, journee, saison", font=small_font, fill="#303030")
    draw.text((1235, 420), "API-Football", font=box_font, fill="#1f1f1f")
    draw.text((1215, 456), "Top scorers et enrichissement", font=small_font, fill="#303030")

    draw_arrow(draw, (470, 470), (540, 650), "JDBC / SQL", small_font, dashed=False, text_offset=-18, number=None, color="#2e7d32")
    draw_arrow(draw, (470, 275), (1070, 250), "HTTP / JSON - GET standings", small_font, dashed=False, text_offset=-26, number=None, color="#5b4db1")
    draw_arrow(draw, (470, 405), (1070, 430), "HTTP / JSON - scorers / fallback", small_font, dashed=False, text_offset=-26, number=None, color="#5b4db1")

    footer_y = 880
    for bullet in PHYSICAL_BULLETS:
        draw.text((120, footer_y), f"- {bullet}", font=small_font, fill="#303030")
        footer_y += 28

    image.save(path)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 11) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def add_highlighted_gherkin(cell, steps: Iterable[str]) -> None:
    cell.text = ""
    keywords = ("Etant donne", "Quand", "Alors")
    for index, step in enumerate(steps):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        for keyword in keywords:
            if step.startswith(keyword):
                keyword_run = paragraph.add_run(keyword)
                keyword_run.bold = True
                keyword_run.font.size = Pt(11)
                keyword_run.font.highlight_color = 7  # yellow
                rest = step[len(keyword):]
                if rest:
                    rest_run = paragraph.add_run(rest)
                    rest_run.font.size = Pt(11)
                break
        else:
            run = paragraph.add_run(step)
            run.font.size = Pt(11)


def style_table_borders(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.first_child_found_in("w:tcBorders")
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for edge in ("top", "left", "bottom", "right"):
                element = borders.find(qn(f"w:{edge}"))
                if element is None:
                    element = OxmlElement(f"w:{edge}")
                    borders.append(element)
                element.set(qn("w:val"), "single")
                element.set(qn("w:sz"), "8")
                element.set(qn("w:color"), "000000")


def set_page_margins(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.905)
    section.bottom_margin = Cm(1.905)
    section.left_margin = Cm(1.905)
    section.right_margin = Cm(1.905)


def add_title(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(32)


def add_section_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(13)


def add_backlog_table(document: Document) -> None:
    add_section_heading(document, "1. Sprint Backlog  Sprint 2")
    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["ID", "User Story", "Priorite", "Points", "Taches", "Responsable", "Statut"]
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(header_cells[idx], header, bold=True)
        set_cell_shading(header_cells[idx], "D9D9D9")

    for item in SPRINT_BACKLOG:
        row_cells = table.add_row().cells
        values = [
            item["id"],
            item["story"],
            item["priority"],
            item["points"],
            item["tasks"],
            item["owner"],
            item["status"],
        ]
        for idx, value in enumerate(values):
            set_cell_text(row_cells[idx], value, bold=False)

    style_table_borders(table)
    document.add_paragraph("")


def add_story_tests(document: Document) -> None:
    add_section_heading(document, "2. Tableau des Story Tests")

    intro = document.add_paragraph()
    run = intro.add_run("User Story testee : ")
    run.bold = True
    intro.add_run("US-TM-02  Consulter le classement detaille d'un championnat")

    table = document.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["User Story", "Story test d'acceptation", "Story Test de refus"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "D9D9D9")

    story_cell = table.rows[1].cells[0]
    story_cell.text = ""
    p = story_cell.paragraphs[0]
    run = p.add_run(STORY_UNDER_TEST)
    run.bold = True
    run.italic = True
    run.font.size = Pt(11)

    add_highlighted_gherkin(table.rows[1].cells[1], STORY_ACCEPTANCE)
    add_highlighted_gherkin(table.rows[1].cells[2], STORY_REJECTION)

    style_table_borders(table)
    document.add_paragraph("")


def add_image_with_caption(document: Document, image_path: Path, caption: str, width: float) -> None:
    if caption:
        p = document.add_paragraph()
        r = p.add_run(caption)
        r.bold = True
    document.add_picture(str(image_path), width=Inches(width))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_diagrams(document: Document) -> None:
    add_section_heading(document, "3. Diagramme de Sequence  Consultation du classement d'une competition (US-TM-02)")
    add_image_with_caption(document, SEQUENCE_PNG, "", 7.0)
    document.add_paragraph("")

    add_section_heading(document, "4. Architectures logique et physique")
    add_image_with_caption(document, LOGICAL_PNG, "Architecture logique :", 6.9)
    document.add_paragraph("")
    add_image_with_caption(document, PHYSICAL_PNG, "Architecture physique :", 6.9)


def generate_docx() -> None:
    document = Document()
    set_page_margins(document)
    add_title(document)
    document.add_paragraph("")
    add_backlog_table(document)
    add_story_tests(document)
    add_diagrams(document)
    document.save(OUTPUT_DOCX)


def main() -> None:
    ensure_output_dir()
    create_sequence_diagram(SEQUENCE_PNG)
    draw_layered_architecture(LOGICAL_PNG)
    draw_physical_architecture(PHYSICAL_PNG)
    generate_docx()
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
